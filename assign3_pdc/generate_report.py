#!/usr/bin/env python3
"""
Build a Word (.docx) lab report from ``benchmark_runs.csv`` + ``plot_out/*.png``.

Includes measured timings, throughput, actual vs ideal speedup / parallel efficiency,
and embedded figures from ``visualize.generate_all_plots``.
"""

from __future__ import annotations

import argparse
from datetime import datetime
from pathlib import Path

import pandas as pd


PLOT_FILENAMES = (
    "plot1_execution_time_vs_workers.png",
    "plot2_speedup_vs_workers.png",
    "plot3_pages_per_second_four_configs.png",
    "plot4_worker_utilization_heatmap.png",
)

PLOT_CAPTIONS = (
    "Execution wall time versus worker count for threaded, multiprocessing, and Ray crawlers.",
    "Observed speedup versus worker count with dashed ideal linear reference (baseline = 1 worker).",
    "Throughput (pages/s) bar chart: Sequential vs parallel engines at eight workers.",
    "Worker utilization heat map (approximate cumulative % vs time bucket × worker ID).",
)


def _resolve_plot_path(plot_dir: Path, fname: str, *, csv_dir: Path) -> Path:
    """Prefer ``plot_dir``, then CSV directory (figures often copied next to benchmark CSV)."""
    primary = plot_dir / fname
    if primary.exists():
        return primary
    fallback = csv_dir / fname
    if fallback.exists():
        return fallback
    return primary


def _ensure_plots(*, csv_path: Path, metrics_path: Path, plot_dir: Path) -> None:
    import matplotlib

    matplotlib.use("Agg")

    from visualize import generate_all_plots

    plot_dir.mkdir(parents=True, exist_ok=True)
    generate_all_plots(
        metrics_path=metrics_path,
        benchmark_path=csv_path,
        output_dir=plot_dir,
    )


def _speedup_table(df: pd.DataFrame) -> pd.DataFrame:
    parallel = df[df["crawler"].isin(["threaded", "multiprocessing", "ray"])].copy()

    rows: list[dict[str, float | int | str]] = []
    for crawler in sorted(parallel["crawler"].unique()):
        blk = parallel[parallel["crawler"] == crawler].sort_values("num_workers")
        ref = blk.loc[blk["num_workers"] == 1, "execution_time_sec"]
        if ref.empty:
            continue
        t1 = float(ref.iloc[0])
        for _, r in blk.iterrows():
            w = int(r["num_workers"])
            te = float(r["execution_time_sec"])
            ideal = float(w)
            actual = t1 / max(te, 1e-9)
            eff = (actual / ideal * 100.0) if ideal else 0.0
            rows.append(
                {
                    "crawler": crawler,
                    "workers": w,
                    "time_s": round(te, 2),
                    "actual_speedup": round(actual, 2),
                    "ideal_speedup": int(w),
                    "parallel_efficiency_pct": round(eff, 1),
                }
            )
    return pd.DataFrame(rows)


def _embed_extra_images(doc, paths: list[Path], *, width_inches: float) -> None:
    from docx.shared import Inches

    for pth in paths:
        pth = Path(pth).expanduser().resolve()
        if not pth.is_file():
            doc.add_paragraph(f"[Skipping missing optional image: {pth}]")
            continue
        doc.add_picture(str(pth), width=Inches(width_inches))
        cap = doc.add_paragraph(str(pth.name))
        cap.runs[0].italic = True


def build_word_report(
    *,
    csv_path: Path,
    plot_dir: Path,
    output_docx: Path,
    seed_url: str = "https://books.toscrape.com",
    benchmark_max_pages: int = 200,
    extra_image_paths: list[Path] | None = None,
) -> None:
    csv_dir = csv_path.parent
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Inches, Pt
    except ImportError as exc:  # pragma: no cover
        raise SystemExit(
            "Missing python-docx. Install with: pip install python-docx"
        ) from exc

    df = pd.read_csv(csv_path)
    df.columns = [str(c).strip().lower() for c in df.columns]
    if "crawler" in df.columns:
        df["crawler"] = df["crawler"].astype(str).str.strip()
    df["num_workers"] = pd.to_numeric(df["num_workers"], errors="coerce").fillna(1).astype(int)
    df["execution_time_sec"] = pd.to_numeric(df["execution_time_sec"], errors="coerce")
    df["pages_per_second"] = pd.to_numeric(df["pages_per_second"], errors="coerce")

    worker_vals = sorted({int(w) for w in df["num_workers"].dropna().tolist()})

    doc = Document()

    title = doc.add_heading("Parallel Web Crawlers – Benchmark Report", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(
        f"Generated {datetime.now().strftime('%Y-%m-%d %H:%M')}\n"
        f"[Add your name / course / ID here]"
    )
    run.font.size = Pt(11)

    doc.add_paragraph()

    doc.add_heading("1. Purpose", level=1)
    doc.add_paragraph(
        "This document compares four crawling strategies (sequential baseline, multi-threaded "
        "BFS with load monitoring, multiprocessing with manager-backed shared state, and Ray "
        "actors) using the same seed site. It records wall-clock time and throughput, compares "
        "observed speedups to ideal linear scaling, and attaches the four figures produced by "
        "the course visualization pipeline."
    )

    doc.add_heading("2. Experimental setup", level=1)
    doc.add_paragraph(
        f"• Seed URL used for the assignment runs: {seed_url}\n"
        f"• Target successful page count for full benchmark mode: {benchmark_max_pages}\n"
        f"• Worker counts appearing in this CSV: {worker_vals}\n"
        "  (sequential runs use a single worker by definition)\n"
        "• Host: same machine and network for all rows in the attached CSV (replace this "
        "sentence with your hardware and OS when you submit).\n\n"
        "The tables and charts below are built from the CSV at the time you generated this "
        "document. Replace `benchmark_runs.csv` with output from "
        "`python main.py --mode benchmark` to report your own measured numbers."
    )

    doc.add_heading("3. Benchmark methodology", level=1)
    doc.add_paragraph(
        "The benchmark compares engines under the same crawling task: identical seed URL, shared "
        "URL-normalization rules, and a fixed target successful page count in full benchmark "
        "mode (`main.py --mode benchmark`). Reported **`execution_time_sec`** is wall-clock "
        "time for the crawl to finish (or stop at the configured cap), not CPU time summed "
        "across cores. **`pages_per_second`** is total successful pages divided by that "
        "wall-clock duration for the run—the standard throughput snapshot for comparing "
        "strategies."
    )
    doc.add_paragraph(
        "Each parallel engine also runs multiple **worker-count** settings (typically 1, 2, 4, "
        "and higher) so curves show how latency and throughput evolve as concurrency grows. "
        "This is effectively **fixed problem size** scaling: the work (page budget) stays the "
        "same while hardware parallelism increases. Results depend on disk, DNS, WAN latency, "
        "and polite rate limits if enabled; rerun on your machine before submitting definitive "
        "numbers."
    )

    doc.add_heading("4. Design trade-offs across engines", level=1)
    doc.add_paragraph(
        "**Sequential crawler.** Simple control flow and minimal synchronization: easy to "
        "reason about correctness and backlog ordering, predictable memory use. Throughput "
        "is capped by doing one outstanding fetch pipeline at a time; it establishes a baseline "
        "latency for the chosen page budget."
    )
    doc.add_paragraph(
        "**Threaded crawler (shared-memory BFS, work stealing, monitor).** Multiple worker "
        "threads share queues and visited structures protected by locks. For I/O-bound HTTP "
        "work threads often overlap network waits efficiently; cost shows up as lock "
        "contention, queue management, and any serial sections that must run on one thread. "
        "The utilization heat map (Figure 4) summarizes how unevenly threads finish batches "
        "over time."
    )
    doc.add_paragraph(
        "**Multiprocessing crawler (manager-backed frontier).** Processes get true parallel "
        "Python interpreters and bypass the global interpreter lock between processes, but pay "
        "for **IPC**: serializing frontier updates through a manager introduces latency and "
        "can bottleneck at high churn. Fits CPU-ish post-processing or isolation needs; pure "
        "fetch-heavy workloads may still be network-bound."
    )
    doc.add_paragraph(
        "**Ray actors (master / workers).** Task placement and queues are handled by Ray’s "
        "runtime: good for structuring distributed-style coordination inside one cluster or "
        "one machine. There is serialized messaging and scheduler overhead; Ray also has a "
        "noticeable startup cost unless the runtime is warmed. Prefer when you explicitly "
        "want durable actor semantics and integration with Ray’s ecosystem."
    )

    doc.add_heading("5. Raw timing results", level=1)
    doc.add_paragraph(
        "Each row is one configuration. Throughput is inferred as pages divided by wall time "
        "for that run (as stored in the CSV)."
    )

    t1 = doc.add_table(rows=1, cols=len(df.columns))
    t1.style = "Table Grid"
    for j, col in enumerate(df.columns):
        t1.rows[0].cells[j].text = str(col)
    for _, row in df.iterrows():
        cells = t1.add_row().cells
        for j, col in enumerate(df.columns):
            cells[j].text = str(row[col])

    doc.add_paragraph()

    doc.add_heading("6. Actual speedup vs ideal (parallel engines)", level=1)
    doc.add_paragraph(
        "Ideal (linear) speedup assumes the 1-worker runtime would shrink by a factor of "
        "N when using N perfectly independent workers—rare for I/O-heavy Python crawls because "
        "of coordination overhead, network limits, and contention. "
        "Actual speedup uses the measured 1-worker time T₁ for the same engine: T₁ / Tₙ. "
        "Parallel efficiency is actual speedup divided by N."
    )

    sp = _speedup_table(df)
    if not sp.empty:
        tbl = doc.add_table(rows=1, cols=len(sp.columns))
        tbl.style = "Table Grid"
        for j, col in enumerate(sp.columns):
            tbl.rows[0].cells[j].text = str(col)
        for _, row in sp.iterrows():
            cc = tbl.add_row().cells
            for j, col in enumerate(sp.columns):
                cc[j].text = str(row[col])
    else:
        doc.add_paragraph("(No parallel rows found in CSV.)")

    doc.add_paragraph()

    doc.add_heading("7. Figures (Plots 1–4)", level=1)
    doc.add_paragraph(
        "These are the four coursework charts (from `visualize.generate_all_plots`, 150 DPI). "
        "They are printed here **before** the reading guide so you see the graphics next to "
        "the tables above. Regenerate with `python generate_report.py --refresh-plots`. "
        f"Plots are resolved from `{plot_dir.name}/` first, then the same folder as this CSV "
        "if filenames match."
    )

    if extra_image_paths:
        doc.add_paragraph(
            "Optional extra figures (your own diagrams or exports) precede the standard plots "
            "when provided via `--extra-image`."
        )
        _embed_extra_images(doc, extra_image_paths, width_inches=6.0)

    pic_width = Inches(6.0)
    for fname, caption in zip(PLOT_FILENAMES, PLOT_CAPTIONS):
        path = _resolve_plot_path(plot_dir, fname, csv_dir=csv_dir)
        doc.add_paragraph()
        if path.exists():
            doc.add_picture(str(path), width=pic_width)
            cap = doc.add_paragraph(caption)
            cap.runs[0].italic = True
        else:
            doc.add_paragraph(
                f"[Missing figure `{fname}` — run with `--plots-dir .` pointing at the PNGs "
                f"or use `--refresh-plots`. Expected `{path}`]"
            )

    doc.add_heading("8. How to read the figures (scalability analysis)", level=1)
    doc.add_paragraph(
        "**Figure 1 (time vs workers).** If doubling workers roughly halves wall time over "
        "some range, the engine exhibits strong concurrency gains on that regime. Flattening "
        "tails mean marginal benefit: contention, remote-server limits, or coordination dominates."
    )
    doc.add_paragraph(
        "**Figure 2 (speedup vs workers).** The dashed diagonal is **ideal linear "
        "speedup**: perfect partitioning with no overhead. Typical crawler curves sit below "
        "it. The gap reflects **Amdahl-style** limits (serial parsing, scheduling, shared "
        "structures) plus **external** limits (HTTP/TCP, bandwidth, site latency). Plateaus "
        "mean concurrency stopped helping."
    )
    doc.add_paragraph(
        "**Figure 3 (throughput bars).** Which backend achieves the highest pages per second "
        "versus sequential at the chosen worker snapshot."
    )
    doc.add_paragraph(
        "**Figure 4 (utilization heatmap).** Rows are coarse time buckets, columns worker "
        "IDs; brighter cells indicate more credited work there. Uneven bands suggest imbalance "
        "or staggered finishes."
    )
    doc.add_paragraph(
        "Use the raw timing and speedup tables (sections 5–6), the checklist in section 9, "
        "and curves above together to quantify the gap versus ideal scaling."
    )

    doc.add_heading("9. Comparison summary (narrative checklist)", level=1)
    max_w = int(df["num_workers"].max()) if not df.empty else 8

    seq = df[df["crawler"] == "sequential"]
    th = df[df["crawler"] == "threaded"]
    threaded_best = th.loc[th["pages_per_second"].idxmax()] if not th.empty else None
    threaded8 = df[(df["crawler"] == "threaded") & (df["num_workers"] == min(8, max_w))]
    seq_pps = float(seq["pages_per_second"].iloc[0]) if not seq.empty else None
    th8_pps = float(threaded8["pages_per_second"].iloc[0]) if not threaded8.empty else None

    bullets = [
        "Measured times are always slower than the hypothetical “ideal” curve because of "
        "Python scheduling, lock contention, serialization in multiprocessing, and Ray "
        "scheduling overhead—not all work parallelizes.",
        "If your real measured speedup stalls around 4–8× even with 16 workers, cite "
        "Amdahl’s law and WAN latency as likely causes.",
    ]
    if seq_pps is not None and threaded_best is not None:
        tb_pps = float(threaded_best["pages_per_second"])
        tw = int(threaded_best["num_workers"])
        ratio = tb_pps / max(seq_pps, 1e-9)
        bullets.append(
            f"CSV snapshot: sequential throughput ≈ {seq_pps:.2f} pages/s vs best threaded "
            f"({tw} workers) ≈ {tb_pps:.2f} pages/s (ratio {ratio:.2f}×)."
        )

    if seq_pps is not None and th8_pps is not None:
        ratio8 = th8_pps / max(seq_pps, 1e-9)
        bullets.append(
            f"If you also compare at eight workers: threaded-8 throughput ≈ {th8_pps:.2f} "
            f"pages/s vs sequential (ratio {ratio8:.2f}×)."
        )

    for line in bullets:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(line)

    doc.add_heading("10. Conclusions", level=1)
    doc.add_paragraph(
        "Summarize here in your own words: which engine scaled best on your hardware, "
        "where diminishing returns appeared, and how close your curves came to ideal speedup "
        "(Figure 2). Mention any outliers (Ray cold start, network throttling)."
    )

    output_docx.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_docx))
    print(f"Wrote {output_docx.resolve()}")


def main() -> None:
    parser = argparse.ArgumentParser(description="Generate Word benchmark report (.docx).")
    parser.add_argument(
        "--csv",
        type=Path,
        default=Path("benchmark_runs.csv"),
        help="benchmark_runs CSV from main.py --mode benchmark",
    )
    parser.add_argument(
        "--plots-dir",
        type=Path,
        default=Path("plot_out"),
        help="Directory containing PNG figures",
    )
    parser.add_argument(
        "--metrics",
        type=Path,
        default=Path("metrics.csv"),
        help="metrics.csv used for heat map (generated by threaded crawler / monitor)",
    )
    parser.add_argument(
        "-o",
        "--output",
        type=Path,
        default=Path("Crawler_Benchmark_Report.docx"),
        help="Output .docx path",
    )
    parser.add_argument(
        "--refresh-plots",
        action="store_true",
        help="Regenerate PNGs via visualize before building the document",
    )
    parser.add_argument(
        "--extra-image",
        type=Path,
        action="append",
        default=[],
        metavar="PATH",
        help="Optional PNG/JPEG to embed after section 9 intro (repeatable)",
    )
    args = parser.parse_args()

    if args.refresh_plots or not any((args.plots_dir / n).exists() for n in PLOT_FILENAMES):
        if not args.refresh_plots:
            print("[generate_report] Some plots missing; refreshing figures…")
        _ensure_plots(
            csv_path=args.csv.resolve(),
            metrics_path=args.metrics.resolve(),
            plot_dir=args.plots_dir.resolve(),
        )

    build_word_report(
        csv_path=args.csv.resolve(),
        plot_dir=args.plots_dir.resolve(),
        output_docx=args.output.resolve(),
        extra_image_paths=[p.resolve() for p in args.extra_image],
    )


if __name__ == "__main__":
    main()
